"""清理 module interface 的测试：内存构造文档，不碰磁盘、不碰 GUI。"""
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from cleaner import StyleCategory, clean_document


def make_doc():
    return Document()


def test_unused_custom_paragraph_style_is_deleted():
    doc = make_doc()
    doc.styles.add_style('我的段落样式', WD_STYLE_TYPE.PARAGRAPH)

    deleted = clean_document(doc)

    mine = [s for s in deleted if s.name == '我的段落样式']
    assert len(mine) == 1
    assert mine[0].category is StyleCategory.PARAGRAPH
    assert '我的段落样式' not in [s.name for s in doc.styles]


def test_style_used_in_body_paragraph_is_kept():
    doc = make_doc()
    doc.styles.add_style('正文用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph('内容', style='正文用到的样式')

    deleted = clean_document(doc)

    assert '正文用到的样式' not in [s.name for s in deleted]
    assert '正文用到的样式' in [s.name for s in doc.styles]


def test_builtin_style_never_deleted_even_when_unused():
    doc = make_doc()
    # 'Heading 1' 是内置样式，默认模板中未被使用
    assert 'Heading 1' in [s.name for s in doc.styles]

    deleted = clean_document(doc)

    assert 'Heading 1' not in [s.name for s in deleted]
    assert 'Heading 1' in [s.name for s in doc.styles]


def test_character_style_used_by_run_is_kept():
    doc = make_doc()
    doc.styles.add_style('在用的字符样式', WD_STYLE_TYPE.CHARACTER)
    doc.styles.add_style('没用的字符样式', WD_STYLE_TYPE.CHARACTER)
    p = doc.add_paragraph('内容')
    run = p.add_run('片段')
    run.style = doc.styles['在用的字符样式']

    deleted = clean_document(doc)

    assert '在用的字符样式' in [s.name for s in doc.styles]
    mine = [s for s in deleted if s.name == '没用的字符样式']
    assert len(mine) == 1
    assert mine[0].category is StyleCategory.CHARACTER


def test_table_style_used_by_table_is_kept():
    doc = make_doc()
    doc.styles.add_style('在用的表格样式', WD_STYLE_TYPE.TABLE)
    doc.styles.add_style('没用的表格样式', WD_STYLE_TYPE.TABLE)
    table = doc.add_table(rows=1, cols=1)
    table.style = doc.styles['在用的表格样式']

    deleted = clean_document(doc)

    assert '在用的表格样式' in [s.name for s in doc.styles]
    mine = [s for s in deleted if s.name == '没用的表格样式']
    assert len(mine) == 1
    assert mine[0].category is StyleCategory.TABLE


def test_style_used_inside_table_cell_is_kept():
    doc = make_doc()
    doc.styles.add_style('表格内段落样式', WD_STYLE_TYPE.PARAGRAPH)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].style = doc.styles['表格内段落样式']

    deleted = clean_document(doc)

    assert '表格内段落样式' not in [s.name for s in deleted]
    assert '表格内段落样式' in [s.name for s in doc.styles]


# ---- 全位置扫描（#4）：页眉/页脚、脚注/尾注、文本框中的使用不算"未使用" ----

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_V_NS = 'urn:schemas-microsoft-com:vml'
_NOTE_RELTYPE_BASE = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/'
)


def _add_note_part(doc, kind, inner_xml):
    """挂一个脚注/尾注 part——python-docx 未暴露 API，只能从 OPC 包层构造。"""
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    wrapper = f'{kind}s'  # footnotes / endnotes
    blob = (
        f'<w:{wrapper} xmlns:w="{_W_NS}">'
        f'<w:{kind} w:id="2">{inner_xml}</w:{kind}>'
        f'</w:{wrapper}>'
    ).encode('utf-8')
    part = Part(
        PackURI(f'/word/{wrapper}.xml'),
        'application/xml',
        blob,
        doc.part.package,
    )
    doc.part.relate_to(part, _NOTE_RELTYPE_BASE + wrapper)


def _add_textbox_with_style(doc, style_id):
    """在正文末尾放一个文本框，段落样式只在框内使用。"""
    from docx.oxml import parse_xml

    paragraph = parse_xml(
        f'<w:p xmlns:w="{_W_NS}" xmlns:v="{_V_NS}">'
        '<w:r><w:pict><v:shape><v:textbox><w:txbxContent>'
        '<w:p><w:pPr>'
        f'<w:pStyle w:val="{style_id}"/>'
        '</w:pPr><w:r><w:t>文本框内容</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>'
    )
    doc.element.body.append(paragraph)


def test_style_used_only_in_header_is_kept():
    doc = make_doc()
    doc.styles.add_style('页眉用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    header = doc.sections[0].header
    header.paragraphs[0].style = doc.styles['页眉用到的样式']

    deleted = clean_document(doc)

    assert '页眉用到的样式' not in [s.name for s in deleted]
    assert '页眉用到的样式' in [s.name for s in doc.styles]


def test_style_used_only_in_footer_is_kept():
    doc = make_doc()
    doc.styles.add_style('页脚用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    footer = doc.sections[0].footer
    footer.paragraphs[0].style = doc.styles['页脚用到的样式']

    deleted = clean_document(doc)

    assert '页脚用到的样式' not in [s.name for s in deleted]
    assert '页脚用到的样式' in [s.name for s in doc.styles]


def test_style_used_only_in_footnote_is_kept():
    doc = make_doc()
    style = doc.styles.add_style('脚注用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    _add_note_part(doc, 'footnote',
                   '<w:p><w:pPr>'
                   f'<w:pStyle w:val="{style.style_id}"/>'
                   '</w:pPr><w:r><w:t>注释内容</w:t></w:r></w:p>')

    deleted = clean_document(doc)

    assert '脚注用到的样式' not in [s.name for s in deleted]
    assert '脚注用到的样式' in [s.name for s in doc.styles]


def test_style_used_only_in_endnote_is_kept():
    doc = make_doc()
    style = doc.styles.add_style('尾注用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    _add_note_part(doc, 'endnote',
                   '<w:p><w:pPr>'
                   f'<w:pStyle w:val="{style.style_id}"/>'
                   '</w:pPr><w:r><w:t>注释内容</w:t></w:r></w:p>')

    deleted = clean_document(doc)

    assert '尾注用到的样式' not in [s.name for s in deleted]
    assert '尾注用到的样式' in [s.name for s in doc.styles]


def test_table_style_used_only_in_footnote_is_kept():
    doc = make_doc()
    style = doc.styles.add_style('脚注表格样式', WD_STYLE_TYPE.TABLE)
    _add_note_part(doc, 'footnote',
                   '<w:tbl><w:tblPr>'
                   f'<w:tblStyle w:val="{style.style_id}"/>'
                   '</w:tblPr><w:tr><w:tc><w:p/></w:tc></w:tr></w:tbl>')

    deleted = clean_document(doc)

    assert '脚注表格样式' not in [s.name for s in deleted]
    assert '脚注表格样式' in [s.name for s in doc.styles]


def test_character_style_used_only_in_header_run_is_kept():
    doc = make_doc()
    doc.styles.add_style('页眉字符样式', WD_STYLE_TYPE.CHARACTER)
    header = doc.sections[0].header
    run = header.paragraphs[0].add_run('页眉内容')
    run.style = doc.styles['页眉字符样式']

    deleted = clean_document(doc)

    assert '页眉字符样式' not in [s.name for s in deleted]
    assert '页眉字符样式' in [s.name for s in doc.styles]


def test_style_used_only_in_textbox_is_kept():
    doc = make_doc()
    style = doc.styles.add_style('文本框用到的样式', WD_STYLE_TYPE.PARAGRAPH)
    _add_textbox_with_style(doc, style.style_id)

    deleted = clean_document(doc)

    assert '文本框用到的样式' not in [s.name for s in deleted]
    assert '文本框用到的样式' in [s.name for s in doc.styles]
